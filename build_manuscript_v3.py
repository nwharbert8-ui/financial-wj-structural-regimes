"""
PhysicaA Manuscript V3 - Language softening + compression + highlights cleanup.
Applied changes from v2:
  1. "mechanistically" -> "structurally" / "statistically"
  2. Killed "strongest reported"; replaced with "strong" / accurate scope
  3. Compressed intro contribution list 7 -> 4
  4. Merged Methods 2.6/2.7; folded 2.8 into 2.9
  5. Discussion 4.1-4.5 compressed to 4.1-4.3
  6. Highlights rephrased for clarity (2 of 5 lines)
"""
import os
import shutil
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT = r"G:\My Drive\inner_architecture_research\financial_wj_fundamental"
OUT_DIR = os.path.join(PROJECT, "PhysicaA_Submission_v3")
os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(os.path.join(OUT_DIR, "Main_Figures"), exist_ok=True)
os.makedirs(os.path.join(OUT_DIR, "Main_Tables"), exist_ok=True)


def setup(doc):
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)


def add_para(doc, text="", bold=False, italic=False, size=11,
             alignment=None, space_before=0, space_after=8,
             line_spacing=1.5, indent_first=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        p.alignment = alignment
    if indent_first:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if text:
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
    return p


def add_heading(doc, text, level=1, before=14, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(13 if level == 1 else 12)
    r.bold = True
    return p


def add_word_table(doc, rows):
    if not rows:
        return
    n_rows = len(rows)
    n_cols = len(rows[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            r.font.name = "Times New Roman"
            r.font.size = Pt(9)
            if i == 0:
                r.bold = True


def table_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.bold = True


# ============================================================================
# MANUSCRIPT
# ============================================================================
print("Building Manuscript_v3.docx...")
doc = Document()
setup(doc)

# Title
add_para(doc,
    "A Pairing-Family Decomposition of the Weighted Jaccard Index Reveals "
    "Direction-Specific Reorganization Modes and Macro-Stress-Coupled "
    "Sign-Inversion Dynamics in S&P 500 Correlation Networks (2003-2025)",
    bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

add_para(doc, "Drake H. Harbert",
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Inner Architecture LLC, Canton, OH 44721, United States",
         size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "ORCID: 0009-0007-7740-3616",
         size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, "Correspondence: Drake@innerarchitecturellc.com",
         size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

# Highlights (revised — line 1 cuts "-driven"; line 2 reworded;
# line 4 clarified; line 5 swaps "sign-driven" for "sign-component")
add_heading(doc, "Highlights")
highlights = [
    "WJ pairing-family decomposition separates magnitude and sign components",
    "Sign-inversion fraction tracks realized volatility (rho=-0.87) across 204 windows",
    "Convergence and divergence regimes occupy distinct pairing-family quadrants",
    "Combined detectors identify six unsigned and one signed-only regime 2003-2025",
    "QE-era is deepest unsigned event; 2022 visible only through sign-component detection",
]
for h in highlights:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(h)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

# Abstract (revised: "mechanistically" -> "structurally"; killed "strongest
# reported"; softened "dominantly *-driven" to "characterized by")
add_heading(doc, "Abstract")
add_para(doc,
    "Existing frameworks for financial correlation network analysis treat "
    "reorganization as unidirectional: episodes are detected when "
    "correlations converge upward during crises. The Weighted Jaccard "
    "(WJ) similarity index extends to bidirectional detection, but a "
    "single WJ value collapses two structurally distinct modes of "
    "reorganization—magnitude reorganization and sign-polarity "
    "reorganization—into one number. This paper applies a pairing-family "
    "decomposition of WJ to disentangle these two modes across 478 "
    "individual S&P 500 constituent stocks (114,003 pairwise Spearman "
    "correlations, 2003-2025).",
    indent_first=True)
add_para(doc,
    "Two pairings are computed at every rolling 252-day window (n=264 "
    "windows) and at six algorithmically detected structural regimes: a "
    "Type 1 continuous-discrete pairing (unsigned WJ minus binary Jaccard "
    "at the top 5% correlation threshold) and a Type 2 sign-treatment "
    "pairing (signed WJ minus unsigned WJ). The decomposition reveals "
    "that convergence and divergence regimes occupy distinct quadrants "
    "of pairing space: convergence regimes show large Type 1 gaps "
    "(mean 0.595) and modest Type 2 gaps (mean 0.177); divergence "
    "regimes show the opposite (mean 0.327 and 0.227 respectively). "
    "Sign-inversion percentage—the fraction of reorganization "
    "attributable to sign change rather than magnitude change—rank-"
    "correlates with mean realized volatility at rho = -0.873 "
    "(p < 10^-64) and with mean VIX at rho = -0.857 (p < 10^-59) across "
    "204 valid windows, a strong macro-stress signature for equity "
    "correlation networks. Crisis-period reorganization is characterized "
    "by magnitude shifts; calm-period and QE-era reorganization is "
    "characterized by sign shifts.",
    indent_first=True)
add_para(doc,
    "A parallel signed-WJ rolling trajectory detects five regimes "
    "(all p < 10^-3 by 5,000-permutation testing), four overlapping with "
    "the unsigned detection (GFC, QE-era, two COVID-adjacent episodes) "
    "and one previously unidentified 2022 regime (2022-07-13 to "
    "2022-08-11; signed WJ = 0.923, z = -13.20) invisible to the "
    "unsigned trajectory because the 2022 reorganization was "
    "predominantly sign-component. The combined catalog spans seven "
    "unique structural regimes 2003-2025. Frobenius distance, an "
    "alternative matrix similarity metric, fails to distinguish the six "
    "unsigned regimes (uniform 0.090-0.098), establishing that the WJ "
    "pairing family captures structural information invisible to global "
    "matrix-distance comparison. The QE-era 2017 decorrelation remains "
    "the deepest unsigned event but is now identified as structurally "
    "distinct from the Global Financial Crisis: comparable depth in "
    "unsigned terms, opposite signature in pairing-family terms.",
    indent_first=True)
add_para(doc,
    "These findings reframe financial correlation regime detection as the "
    "characterization of two distinct reorganization modes—magnitude-"
    "component and sign-component—each with opposite stress dependencies. "
    "The pairing-family decomposition recovers a structural information "
    "channel invisible to unsigned WJ alone.",
    indent_first=True)
add_para(doc,
    "Keywords: weighted Jaccard; pairing-family decomposition; "
    "correlation networks; sign-inversion; structural regime detection; "
    "S&P 500; macro-stress signature; Spearman rank correlation",
    italic=True, space_after=6)
add_para(doc, "JEL Classification: G01, G11, G12, C38, C58",
         italic=True, space_after=14)

# 1. Introduction (compressed contribution list 7 -> 4)
add_heading(doc, "1 Introduction")
add_para(doc,
    "Financial crises are characterized by systematic reorganization of "
    "inter-asset correlations. Portfolio diversification—the foundation "
    "of modern asset allocation since Markowitz [1]—fails when assets "
    "that appear uncorrelated during stable periods become correlated "
    "during stress. A large literature has developed tools to detect and "
    "characterize this reorganization: Dynamic Conditional Correlation "
    "[2], variance-decomposition spillover networks [3], absorption "
    "ratios [4], and Granger-causality connectedness [5]. Network "
    "methods including minimum spanning trees [6], random matrix "
    "filtering [7,8], and hierarchical filtering [9] provide structural "
    "representations of correlation matrices.",
    indent_first=True)
add_para(doc,
    "A shared assumption underlies this literature: correlation "
    "reorganization is unidirectional. Detection frameworks are "
    "calibrated to identify episodes where correlations increase—the "
    "convergence toward unity that characterizes financial crises. The "
    "possibility that correlations may reorganize in the opposite "
    "direction—systematic decorrelation producing equally dramatic "
    "structural change—has not been treated as a detection target. "
    "Periods of historically low correlation, such as the quantitative "
    "easing (QE) era of 2016-2017, are classified as 'stable' or "
    "'low-volatility' rather than as structural regimes in their own "
    "right.",
    indent_first=True)
add_para(doc,
    "A unified framework based on the weighted Jaccard (WJ) similarity "
    "index applied to Spearman rank correlations among individual "
    "S&P 500 constituent stocks eliminates the directional assumption: "
    "WJ measures the proportional overlap between two correlation "
    "vectors, detecting reorganization regardless of direction. The WJ "
    "framework has been applied to industrial sensor networks [16] and "
    "ecological systems [17], where it detects structural reorganization "
    "preceding mechanical failure and trophic state transitions "
    "respectively. Direct application to financial correlation networks "
    "identifies six statistically significant structural regimes "
    "2003-2025, including the QE-era decorrelation as the deepest "
    "unsigned reorganization in the dataset.",
    indent_first=True)
add_para(doc,
    "However, a single WJ value—even when computed bidirectionally—"
    "collapses two structurally distinct components of reorganization "
    "into one number. Correlations between two stocks can change in two "
    "fundamentally different ways: their absolute magnitude can shift "
    "(both stocks become more or less correlated together), or their "
    "sign can flip (a positive correlation becomes negative, or vice "
    "versa). The unsigned WJ index registers the absolute magnitude of "
    "correlation change but is blind to the sign component; the signed "
    "WJ index registers both. The gap between them is the sign-flip "
    "share of reorganization. This paper applies a pairing-family "
    "decomposition of WJ [16] to disentangle these two components: a "
    "Type 1 continuous-discrete pairing (unsigned WJ minus binary "
    "Jaccard at the top 5% correlation threshold), which separates "
    "bulk-distribution from tail-concentrated reorganization; and a "
    "Type 2 sign-treatment pairing (signed WJ minus unsigned WJ), which "
    "separates magnitude from sign reorganization.",
    indent_first=True)
add_para(doc,
    "The paper makes four empirical contributions. First, an unsigned "
    "WJ rolling-trajectory framework identifies six structural regimes "
    "(three convergence, three divergence) including the QE-era "
    "decorrelation as the deepest unsigned event. Second, the "
    "pairing-family decomposition assigns convergence and divergence "
    "regimes to distinct quadrants of pairing space, revealing two "
    "structurally distinct reorganization modes (magnitude-component "
    "and sign-component). Third, sign-inversion percentage rank-"
    "correlates with macro stress at rho = -0.85 to -0.87 (p < 10^-59) "
    "across 204 valid windows. Fourth, a parallel signed-WJ trajectory "
    "detects a 2022 structural regime invisible to unsigned detection, "
    "extending the regime catalog to seven unique regimes 2003-2025.",
    indent_first=True)

# 2. Methods (merged 2.6/2.7; folded 2.8)
add_heading(doc, "2 Data and Methodology")

add_heading(doc, "2.1 Data Sources", level=2)
add_para(doc,
    "Historical S&P 500 constituent membership is obtained from the "
    "fja05680/sp500 repository (MIT license), providing point-in-time "
    "index composition from 1996 through January 2026. Daily adjusted "
    "closing prices are obtained from Yahoo Finance. Quality filtering "
    "retains stocks with at least 80% trading-day coverage and non-zero "
    "return variance, yielding a final universe of 478 stocks, 5,785 "
    "daily log-return observations (2 January 2003 through 30 December "
    "2025), and 114,003 unique pairwise correlations. GICS sector "
    "classifications are retained exclusively for post-discovery "
    "comparison. CBOE Volatility Index (VIX) historical daily data is "
    "obtained from Yahoo Finance (2003-2025; n = 5,786 trading days).",
    indent_first=True)

add_heading(doc, "2.2 Spearman Correlations and the Weighted Jaccard Index",
            level=2)
add_para(doc,
    "Pairwise Spearman rank correlations are computed for all 114,003 "
    "pairs within each temporal window. Spearman is used as the primary "
    "correlation method because it is robust to the non-normal return "
    "distributions that characterize equity markets [10]; Pearson is "
    "computed as a sensitivity check. Given two correlation vectors a "
    "and b of length m, the unsigned weighted Jaccard similarity is "
    "defined on absolute correlations as WJ_unsigned = "
    "sum(min(|a_i|, |b_i|)) / sum(max(|a_i|, |b_i|)). The signed "
    "weighted Jaccard similarity is defined on shifted correlations "
    "(range [0, 2]) as WJ_signed = sum(min(a_i + 1, b_i + 1)) / "
    "sum(max(a_i + 1, b_i + 1)). Both metrics return 1.0 for identical "
    "networks; values below 1.0 indicate reorganization. The unsigned "
    "WJ measures magnitude reorganization and is blind to sign "
    "inversions; the signed WJ measures combined magnitude and sign "
    "reorganization.",
    indent_first=True)

add_heading(doc, "2.3 Pairing-Family Decomposition", level=2)
add_para(doc,
    "Two pairings of the WJ family are computed [16]. The Type 1 "
    "(continuous-discrete) pairing is WJ_unsigned minus binary Jaccard "
    "at the top 5% absolute correlation threshold, where binary Jaccard "
    "is computed on the indicator vectors I(|a_i| >= t) and I(|b_i| "
    ">= t) and t is the 95th percentile of baseline absolute "
    "correlations. The Type 1 gap measures whether reorganization "
    "concentrates in the bulk of the correlation distribution (large "
    "gap) or in the high-correlation tail (small gap). The Type 2 "
    "(sign-treatment) pairing is WJ_signed minus WJ_unsigned, measuring "
    "the share of reorganization attributable to sign change rather "
    "than magnitude change. A direct interpretation as sign-inversion "
    "percentage is given by gap / (1 - WJ_unsigned), the fraction of "
    "total reorganization that is sign-component.",
    indent_first=True)

add_heading(doc, "2.4 Regime Detection: Sliding Baseline", level=2)
add_para(doc,
    "A 252-trading-day rolling window with 21-day step produces 264 WJ "
    "values across the study period. Each window's correlation vector is "
    "compared to the mean correlation vector of the preceding 60 "
    "windows (approximately five trading years), producing a sliding-"
    "baseline WJ trajectory. This avoids two sources of bias: anchoring "
    "to a fixed historical period (which conflates secular evolution "
    "with regime-specific reorganization) and circular reference (which "
    "occurs when regime data contribute to the reference). Regime "
    "boundaries are identified algorithmically: WJ below the trajectory "
    "mean minus one standard deviation defines a reorganization window; "
    "below two standard deviations defines severe reorganization. "
    "Contiguous reorganization windows define episodes. No manually "
    "defined regime dates are used. Sensitivity analysis across "
    "alternative lookback lengths (40, 60, 80 windows) confirms regime "
    "stability. The same trajectory framework is applied independently "
    "to the unsigned WJ and the signed WJ, producing two regime "
    "catalogs.",
    indent_first=True)

add_heading(doc, "2.5 Directional Decomposition", level=2)
add_para(doc,
    "Each unsigned-WJ-detected regime is classified as convergence or "
    "divergence by comparing the mean absolute Spearman correlation "
    "within the regime to the mean absolute correlation of its sliding "
    "baseline. A positive delta indicates convergence; a negative delta "
    "indicates divergence. This decomposition is applied post-detection: "
    "the WJ trajectory identifies regimes without directional "
    "information, and the direction is determined afterward.",
    indent_first=True)

add_heading(doc, "2.6 WJ-Native Baseline, Per-Window Trajectories, and "
            "Macro-Stress Alignment", level=2)
add_para(doc,
    "Return data from all normal-state windows (WJ above the trajectory "
    "mean) are pooled across the entire study period to construct the "
    "WJ-native baseline correlation matrix. This yields 4,809 "
    "normal-state trading days spanning 2003 through 2025, eliminating "
    "the temporal drift inherent in single-era baselines. Per-window "
    "pairing-family trajectories are computed by applying the Type 1 "
    "and Type 2 pairings at every one of the 264 windows against this "
    "native baseline, producing trajectories of unsigned WJ, signed WJ, "
    "binary Jaccard, Type 1 gap, Type 2 gap, and sign-inversion "
    "percentage. Per-window macro stress measures are aligned to these "
    "trajectories. Realized volatility is computed as the rolling "
    "20-day standard deviation of the cross-sectional mean of S&P 500 "
    "stock returns, annualized by sqrt(252), with mean and peak "
    "computed within each window's daily span. Mean and peak VIX are "
    "computed analogously from CBOE VIX daily closes. Cross-trajectory "
    "correlations are tested by Spearman rank correlation across the "
    "204 windows with valid sliding-baseline data (after the 60-window "
    "warmup).",
    indent_first=True)

add_heading(doc, "2.7 Statistical Inference and Alternative Metric "
            "Comparison", level=2)
add_para(doc,
    "Regime-level significance is assessed via permutation testing with "
    "5,000 permutations. For each permutation, the same number of "
    "trading days as the regime are randomly sampled from the full "
    "study period, Spearman correlations are computed, and WJ "
    "(unsigned and signed) is calculated against the WJ-native "
    "baseline. Bootstrap 95% confidence intervals use 1,000 "
    "subject-level resamples. Cascade stability across regime pairs is "
    "measured by Kendall tau. Random seed 42 is used throughout. As an "
    "alternative metric comparison, normalized Frobenius distance "
    "between two correlation matrices A and B is computed as "
    "sqrt(sum((A_ij - B_ij)^2)) / sqrt(2*n*(n-1)), where n is the "
    "number of stocks; this provides a global matrix similarity "
    "comparison independent of the WJ framework.",
    indent_first=True)

add_heading(doc, "2.8 Fingerprint Reorganization and Clustering", level=2)
add_para(doc,
    "Each stock's correlation fingerprint is its vector of absolute "
    "Spearman correlations with all other 477 stocks. Fingerprint "
    "reorganization is defined as 1 - WJ(f_baseline, f_regime), where "
    "f is the stock's fingerprint vector. WJ-native fingerprint "
    "clustering is performed via Ward's hierarchical clustering on the "
    "478 x 478 fingerprint distance matrix (d = 1 - WJ); optimal "
    "cluster count is selected by silhouette score [12] across "
    "K = 2-24.",
    indent_first=True)

# 3. Results
add_heading(doc, "3 Results")

add_heading(doc, "3.1 Six Unsigned-Detected Structural Regimes 2003-2025",
            level=2)
add_para(doc,
    "The unsigned WJ rolling trajectory identifies six statistically "
    "significant structural regimes (Fig. 1, Table 1). Three are "
    "convergence episodes: the Global Financial Crisis (May 2008-May "
    "2009; WJ = 0.740, z = -18.08), an early-2020 COVID-onset episode "
    "(January 2020; WJ = 0.739, z = -10.17), and a mid-2020 post-crash "
    "recovery episode (May-August 2020; WJ = 0.754, z = -11.68). Three "
    "are divergence episodes: the QE-era decorrelation (December "
    "2016-September 2017; WJ = 0.681, z = -22.39) and two 2024 "
    "episodes (January 2024, WJ = 0.699, z = -13.06; May-August 2024, "
    "WJ = 0.690, z = -16.82). All six are significant at p < 0.001 by "
    "5,000-permutation testing.",
    indent_first=True)

add_heading(doc, "3.2 QE-Era Decorrelation as the Deepest Unsigned Event",
            level=2)
add_para(doc,
    "The QE-era divergence regime produces the lowest WJ value in the "
    "dataset (0.681) and the strongest permutation z-score (-22.39). "
    "During this period, the mean absolute Spearman correlation dropped "
    "to 0.176—the lowest level in 22 years of data—compared to a "
    "baseline of 0.351. The GFC convergence episode, by contrast, "
    "produced WJ = 0.740 with mean absolute correlation rising to "
    "0.493. WJ measures the proportional overlap between correlation "
    "profiles, so when correlations are uniformly low, the proportional "
    "change relative to the already-low baseline is large. The QE-era "
    "decorrelation reorganized the network more completely—in "
    "proportional terms—than the GFC convergence.",
    indent_first=True)

add_heading(doc, "3.3 Convergence-Divergence Asymmetry", level=2)
add_para(doc,
    "The correlation between convergence reorganization and divergence "
    "reorganization across stocks is moderate (Spearman rho = 0.383, "
    "p = 3.74 x 10^-18; Fig. 5). Stocks that reorganize most during "
    "convergence episodes are partially but not fully overlapping with "
    "those that reorganize most during divergence episodes. Several "
    "stocks exhibit strongly direction-specific reorganization: FMCC "
    "and FNMA reorganize primarily during convergence episodes, "
    "consistent with their role as government-sponsored enterprises "
    "whose correlations with the broader market increase sharply during "
    "financial stress. LLY, HUM, and CAH reorganize primarily during "
    "divergence episodes, consistent with the defensive character of "
    "healthcare stocks that decorrelate from the market during low-"
    "volatility regimes. These direction-specific patterns are "
    "invisible to frameworks that aggregate across regime types.",
    indent_first=True)

add_heading(doc, "3.4 Direction-Dependent Cascade Stability", level=2)
add_para(doc,
    "Cascade stability—the consistency of stock reorganization rankings "
    "across episodes—is strongly direction-dependent (Fig. 8, Table 2). "
    "The strongest same-direction pairs are the two COVID-era "
    "convergence episodes (Episodes 3-4: tau = 0.525, p < 10^-65) and "
    "the two 2024 divergence episodes (Episodes 5-6: tau = 0.664, "
    "p < 10^-104). The mean tau across all same-direction pairs is "
    "0.209; the mean across all cross-direction pairs is 0.119. The "
    "market possesses two distinct vulnerability profiles: stocks near "
    "the top of the convergence cascade are not systematically the "
    "same stocks near the top of the divergence cascade.",
    indent_first=True)

add_heading(doc, "3.5 Data-Driven Architecture vs GICS", level=2)
add_para(doc,
    "WJ-native fingerprint clustering identifies K = 2 as the optimal "
    "cluster count (silhouette = 0.298; Fig. 6). The two clusters "
    "contain 319 and 159 stocks respectively. Agreement with the GICS "
    "sector taxonomy is low: ARI = 0.101, NMI = 0.216. GICS captures "
    "approximately one-tenth of the correlation architecture revealed "
    "by WJ fingerprint similarity. Both clusters are divergence-"
    "dominant (87.5% and 74.2% of stocks reorganize more during "
    "divergence than convergence). The binary split reflects a "
    "structural property of the baseline correlation architecture that "
    "is not reducible to directional response.",
    indent_first=True)

add_heading(doc, "3.6 Extreme Reorganizers Across Both Directions", level=2)
add_para(doc,
    "Gap analysis identifies the two stocks with the highest mean "
    "fingerprint reorganization across all six episodes: CVG "
    "(Convergys Corporation; mean reorganization = 0.633) and EP "
    "(El Paso Corporation; mean reorganization = 0.604). The gap "
    "between EP and the third-ranked stock (FMCC, 0.472) is 0.133—the "
    "largest discontinuity in the distribution (Fig. 3). Permutation "
    "testing confirms that this two-stock group reorganizes "
    "significantly more than expected by chance (z = 8.98, p < 0.001). "
    "Both stocks have extremely low baseline connectivity (CVG: mean "
    "|rho| = 0.030; EP: mean |rho| = 0.041), placing them at the "
    "periphery of the correlation network.",
    indent_first=True)

add_heading(doc, "3.7 Pairing-Family Decomposition Distinguishes "
            "Convergence from Divergence Regimes", level=2)
add_para(doc,
    "Application of the Type 1 continuous-discrete pairing and Type 2 "
    "sign-treatment pairing to the six detected regimes reveals "
    "direction-specific signatures (Fig. 9, Table 4). Convergence "
    "regimes show large Type 1 gaps (mean 0.595, range 0.581-0.603) "
    "and modest Type 2 gaps (mean 0.177, range 0.172-0.180). "
    "Divergence regimes show the opposite: smaller Type 1 gaps "
    "(mean 0.327, range 0.302-0.357) and larger Type 2 gaps "
    "(mean 0.227, range 0.221-0.235). Plotting the six regimes in the "
    "Type 1 x Type 2 plane (Fig. 10) yields complete separation: "
    "convergence regimes occupy the high-Type-1 / low-Type-2 quadrant; "
    "divergence regimes occupy the low-Type-1 / high-Type-2 quadrant.",
    indent_first=True)
add_para(doc,
    "This direction-specific signature has a structural interpretation. "
    "Convergence regimes—financial crises in which correlations rise "
    "toward unity across most pairs—produce uniformly large WJ "
    "departures across the bulk of the correlation distribution but few "
    "sign inversions, because pairs that were positively correlated "
    "remain positive. Divergence regimes—periods of unusually low "
    "correlation such as the QE era—produce smaller WJ departures "
    "across the bulk distribution because absolute correlation "
    "magnitudes shrink uniformly, but they produce more sign inversions "
    "because pairs near zero correlation can flip polarity when "
    "correlations decay further. The two reorganization modes are "
    "structurally distinct.",
    indent_first=True)
add_para(doc,
    "Sign-inversion percentage is direction-dependent (Table 5): "
    "convergence regimes show 68-70%, divergence regimes show 73-74%. "
    "Both regime types involve substantial sign reorganization, but the "
    "sign component is consistently higher under divergence.",
    indent_first=True)

add_heading(doc, "3.8 Sign-Inversion Percentage Rank-Correlates with "
            "Macro Stress", level=2)
add_para(doc,
    "The pairing-family decomposition is extended to all 264 rolling "
    "windows, producing trajectories of unsigned WJ, signed WJ, Type 1 "
    "gap, Type 2 gap, and sign-inversion percentage spanning 2003-2025. "
    "Aligning these trajectories to mean and peak VIX and realized "
    "volatility per window (n = 204 valid windows after the 60-window "
    "sliding-baseline warmup) yields the cross-trajectory correlations "
    "summarized in Table 6 and visualized in Fig. 11.",
    indent_first=True)
add_para(doc,
    "Sign-inversion percentage rank-correlates with mean realized "
    "volatility at rho = -0.873 (p < 10^-64), with mean VIX at "
    "rho = -0.857 (p < 10^-59), with peak VIX at rho = -0.639 "
    "(p < 10^-24), and with peak realized volatility at rho = -0.793 "
    "(p < 10^-44). The Type 1 gap correlates positively with all four "
    "macro stress measures (rho = +0.414 to +0.727; p < 10^-9 in all "
    "cases). The Type 2 gap shows weaker positive correlations "
    "(rho = +0.128 to +0.408).",
    indent_first=True)
add_para(doc,
    "The negative correlation of sign-inversion percentage with macro "
    "stress is the strongest cross-trajectory signal in the analysis. "
    "Its interpretation is direct: when realized volatility is high, "
    "correlations across the network spike upward in unison, producing "
    "large unsigned WJ departure but few sign inversions; the "
    "reorganization is characterized by magnitude shifts. When realized "
    "volatility is low, the unsigned WJ departure is smaller but a "
    "larger fraction of the reorganization comes from sign changes; the "
    "reorganization is characterized by sign shifts. The sign-inversion "
    "percentage cleanly separates these two modes across 22 years of "
    "data.",
    indent_first=True)

add_heading(doc, "3.9 Signed WJ Trajectory Detects 2022 Regime Invisible "
            "to Unsigned Detection", level=2)
add_para(doc,
    "A parallel rolling-trajectory analysis on signed WJ (computed "
    "against the same 60-window sliding baseline) identifies five "
    "regimes of signed-WJ depression (1-sigma threshold). Four overlap "
    "directly with the unsigned-detected regimes: GFC (2008-07-08 to "
    "2009-05-07), QE era (2017-01-09 to 2017-09-08), COVID onset "
    "(2020-01-10), and COVID recovery (2020-05-12 to 2020-08-11). The "
    "fifth—2022-07-13 to 2022-08-11—does not overlap with any "
    "unsigned-detected regime.",
    indent_first=True)
add_para(doc,
    "Permutation testing confirms statistical significance for all five "
    "signed regimes (Table 7). The 2022 regime produces signed WJ = "
    "0.923, z = -13.20, p < 10^-3. The QE era is the deepest signed-WJ "
    "event (z = -25.20), preserving the QE-era-as-deepest finding. The "
    "2022 regime spans the early summer 2022 inflation-driven bear "
    "market in U.S. equities. Conventional volatility-based regime "
    "detectors flag this period weakly compared to the GFC or COVID "
    "episodes because absolute volatility levels were lower, but the "
    "correlation network underwent substantial sign-component "
    "reorganization that the unsigned WJ misses by construction—sign "
    "changes that preserve absolute correlation magnitude leave "
    "unsigned WJ unchanged. Conversely, the two 2024 unsigned-detected "
    "divergence regimes (Episodes 5 and 6) do not appear in the signed "
    "WJ detection. The combined regime catalog spans seven unique "
    "structural regimes 2003-2025: four detected by both methods, one "
    "(2022) detected by signed only, two (2024) detected by unsigned "
    "only.",
    indent_first=True)

add_heading(doc, "3.10 Frobenius Distance Does Not Distinguish Regimes",
            level=2)
add_para(doc,
    "Computing the normalized Frobenius distance between each regime's "
    "correlation matrix and the WJ-native baseline yields uniform "
    "values across all six unsigned-detected regimes (Table 4): "
    "GFC = 0.097, QE = 0.095, COVID-onset = 0.098, COVID-recovery = "
    "0.090, 2024-Jan = 0.096, 2024-H2 = 0.096. These are statistically "
    "indistinguishable. The six structural regimes—which span an 8.6-"
    "fold range in WJ permutation z-scores—are not separable by "
    "Frobenius distance. The regime detection therefore depends on the "
    "WJ-pairing-family framework specifically; a correlation network "
    "can be globally close to baseline in the Frobenius sense while "
    "having substantially reorganized through pairwise relational "
    "structure that WJ captures.",
    indent_first=True)

add_heading(doc, "3.11 Robustness Checks", level=2)
add_para(doc,
    "Subject-level bootstrap (1,000 resamples) yields tight 95% "
    "confidence intervals on unsigned WJ for each regime (Table 8): GFC "
    "[0.695, 0.779], QE [0.608, 0.737], COVID-onset [0.680, 0.762], "
    "COVID-recovery [0.710, 0.763], 2024-Jan [0.627, 0.726], 2024-H2 "
    "[0.622, 0.729]. None of the six estimates is consistent with the "
    "null (WJ ~= 1.0). Pearson correlations were computed as a "
    "sensitivity check; Pearson WJ values are systematically lower than "
    "Spearman WJ values (delta = 0.022 to 0.120), indicating that "
    "Pearson detects deeper reorganization than Spearman. The relative "
    "ranking of regimes by reorganization depth is preserved under "
    "both correlation methods. Spearman is retained as the primary "
    "metric on the basis of its robustness to non-normal return "
    "distributions [10].",
    indent_first=True)

# 4. Discussion (4.1+4.2 merged; 4.3+4.4 merged; 4.5 -> 4.3; 4.6 -> 4.4)
add_heading(doc, "4 Discussion")

add_heading(doc, "4.1 Pairing-Family Decomposition Recovers Two "
            "Reorganization Modes", level=2)
add_para(doc,
    "The unsigned WJ rolling-trajectory framework eliminates the "
    "directional assumption of conventional correlation regime "
    "detection: convergence and divergence are both reorganization "
    "events, and the QE-era decorrelation produces the deepest WJ-"
    "measurable structural change in 22 years of S&P 500 data, "
    "exceeding the GFC. The pairing-family decomposition extends this "
    "finding by separating two structurally distinct components of "
    "reorganization—magnitude and sign. The Type 1 continuous-discrete "
    "pairing and Type 2 sign-treatment pairing assign each regime a "
    "position in a two-dimensional decomposition space, revealing that "
    "convergence and divergence regimes occupy distinct quadrants. The "
    "trajectory-level finding that sign-inversion percentage rank-"
    "correlates with macro stress at rho = -0.85 to -0.87 across 204 "
    "windows supports a two-mode interpretation: under high macro "
    "stress, the network reorganizes through coherent magnitude "
    "shifts; under low macro stress, the network reorganizes through "
    "sign flips. The sign-inversion percentage cleanly separates these "
    "two modes because it is the fractional share of reorganization "
    "attributable to sign change rather than magnitude change.",
    indent_first=True)

add_heading(doc, "4.2 The QE-Era and 2022 Regimes in Context", level=2)
add_para(doc,
    "The original finding that the QE-era decorrelation produced the "
    "deepest unsigned WJ in 22 years of data is preserved and extended "
    "by the pairing-family decomposition. The QE era is not just deeper "
    "than the GFC in unsigned terms; it is structurally distinct in "
    "decomposition. The QE-era reorganization is characterized by sign "
    "shifts (Type 2 gap = 0.235; sign-inversion % = 73.6%); the GFC "
    "reorganization is characterized by magnitude shifts (Type 2 gap = "
    "0.178; sign-inversion % = 68.5%). The newly-identified 2022 "
    "structural regime (2022-07-13 to 2022-08-11) is invisible to the "
    "unsigned WJ trajectory but produces signed WJ z = -13.20 "
    "(p < 10^-3 by 5,000-permutation testing). This period coincides "
    "with the early summer 2022 inflation-driven bear market in U.S. "
    "equities. The unsigned WJ trajectory's failure to detect this "
    "regime is not a methodological error but a structural consequence: "
    "sign changes that preserve absolute correlation magnitude are "
    "invisible to unsigned WJ by construction. The signed WJ trajectory "
    "recovers them. The 2022 detection illustrates the practical value "
    "of the signed trajectory as a complement to the unsigned detector "
    "and identifies a market state—sign reorganization without "
    "magnitude reorganization—not characterized in the prior "
    "correlation-network literature.",
    indent_first=True)

add_heading(doc, "4.3 Direction-Dependent Cascades and Practical "
            "Implications", level=2)
add_para(doc,
    "The direction-dependence of cascade ordering (tau = 0.209 same-"
    "direction versus tau = 0.119 cross-direction) indicates that the "
    "market possesses two structurally independent reorganization "
    "modes. A risk model calibrated on convergence-era data may "
    "misidentify which stocks are most vulnerable during a divergence "
    "regime. The stocks that provide the most diversification benefit "
    "during normal times may be the most structurally unstable during "
    "divergence, while the stocks that become most correlated during "
    "crises may be structurally stable during divergence. The "
    "pairing-family framework provides a basis for direction-stratified "
    "risk modeling: regimes can be classified by their pairing-family "
    "fingerprint before deployment of a direction-specific risk metric.",
    indent_first=True)

add_heading(doc, "4.4 Limitations", level=2)
add_para(doc,
    "Several limitations warrant discussion. WJ_unsigned operates on "
    "absolute correlations and is therefore insensitive to sign "
    "changes; the present paper addresses this by introducing the "
    "signed WJ trajectory as a complementary detector. The binary "
    "Jaccard threshold for the Type 1 pairing is set at the top 5% of "
    "baseline correlations; sensitivity analysis at alternative "
    "thresholds (top 1%, top 10%) would strengthen the finding. The "
    "trajectory-level correlation between sign-inversion percentage "
    "and realized volatility is computed across overlapping rolling "
    "windows, which violates strict independence; the appropriate "
    "effective sample size is smaller than n = 204, though the "
    "magnitude of the correlation makes this distinction practically "
    "less consequential. The 478-stock universe excludes stocks with "
    "less than 80% trading-day coverage; survivorship bias is "
    "directionally conservative, as excluded stocks would likely "
    "exhibit more extreme reorganization. The two extreme reorganizers "
    "(CVG, EP) are both delisted, limiting practical applicability of "
    "stock-level findings. The K = 2 clustering solution is coarse; "
    "finer structure may exist at lower silhouette values. The 2022 "
    "signed-only regime has been validated by permutation testing "
    "within this dataset but has not been replicated in independent "
    "equity markets.",
    indent_first=True)

# 5. Conclusion
add_heading(doc, "5 Conclusion")
add_para(doc,
    "Correlation network reorganization in equity markets is "
    "bidirectional and decomposes into two structurally distinct "
    "modes—magnitude and sign—each with opposite stress dependencies. "
    "A unified weighted Jaccard framework with pairing-family "
    "decomposition applied to 478 individual S&P 500 stocks identifies "
    "six unsigned-detected structural regimes 2003-2025 plus one "
    "previously unidentified 2022 sign-component regime, for a total of "
    "seven unique regimes. Sign-inversion percentage rank-correlates "
    "with mean realized volatility at rho = -0.873 (p < 10^-64) across "
    "204 valid windows, a strong macro-stress signature for equity "
    "correlation networks. Convergence and divergence regimes occupy "
    "distinct quadrants of pairing-family space, with convergence "
    "characterized by magnitude shifts and divergence characterized by "
    "sign shifts. The QE-era 2017 decorrelation remains the deepest "
    "unsigned event in 22 years of data and is now identified as "
    "structurally distinct from the Global Financial Crisis. Frobenius "
    "distance fails to distinguish the six unsigned regimes, "
    "establishing that the pairing-family decomposition captures "
    "structural information invisible to global matrix-distance "
    "comparison. Reproducible code for all analyses is available at "
    "https://github.com/nwharbert8-ui/financial-wj-structural-regimes.",
    indent_first=True)

# Required Elsevier sections
add_heading(doc, "Declaration of Competing Interest")
add_para(doc, "The author declares no competing interests.",
         indent_first=True)

add_heading(doc, "Acknowledgments")
add_para(doc,
    "No external funding was received. Historical S&P 500 constituent "
    "data were obtained from the fja05680/sp500 repository (MIT "
    "license). Daily price data and CBOE VIX historical data were "
    "obtained from Yahoo Finance.",
    indent_first=True)

add_heading(doc, "Data Availability")
add_para(doc,
    "All data sources are publicly available. The complete analysis "
    "pipeline and reproducible code are available at "
    "https://github.com/nwharbert8-ui/financial-wj-structural-regimes "
    "(https://doi.org/10.5281/zenodo.19025536).",
    indent_first=True)

add_heading(doc, "CRediT Author Statement")
add_para(doc,
    "Drake H. Harbert: Conceptualization, Methodology, Software, Formal "
    "Analysis, Investigation, Data Curation, Writing - Original Draft, "
    "Writing - Review & Editing, Visualization.",
    indent_first=True)

add_heading(doc, "Ethics Statement")
add_para(doc,
    "Not applicable. This study uses publicly available financial "
    "market data only.",
    indent_first=True)

add_heading(doc, "Declaration of Generative AI Use")
add_para(doc,
    "Claude (Anthropic, Claude Opus 4.7) was used as a programming "
    "assistant during pipeline development, manuscript formatting, and "
    "code review. All analytical decisions, methodology design, data "
    "interpretation, and scientific conclusions are solely the work of "
    "the author. The AI tool was not used to generate scientific text, "
    "interpret results, or formulate hypotheses. All code was reviewed "
    "and validated by the author prior to execution.",
    indent_first=True)

# References
add_heading(doc, "References")
references = [
    "[1] H. Markowitz, Portfolio selection, J. Finance 7 (1952) 77-91.",
    "[2] R. Engle, Dynamic conditional correlation: a simple class of "
    "multivariate GARCH models, J. Bus. Econ. Stat. 20 (2002) 339-350.",
    "[3] F.X. Diebold, K. Yilmaz, On the network topology of variance "
    "decompositions: measuring the connectedness of financial firms, "
    "J. Econometrics 182 (2014) 119-134.",
    "[4] M. Kritzman, Y. Li, S. Page, R. Rigobon, Principal components "
    "as a measure of systemic risk, J. Portf. Manag. 37 (2011) 112-126.",
    "[5] M. Billio, M. Getmansky, A.W. Lo, L. Pelizzon, Econometric "
    "measures of connectedness and systemic risk in the finance and "
    "insurance sectors, J. Financ. Econ. 104 (2012) 535-559.",
    "[6] R.N. Mantegna, Hierarchical structure in financial markets, "
    "Eur. Phys. J. B 11 (1999) 193-197.",
    "[7] L. Laloux, P. Cizeau, J.-P. Bouchaud, M. Potters, Noise "
    "dressing of financial correlation matrices, Phys. Rev. Lett. 83 "
    "(1999) 1467-1470.",
    "[8] V. Plerou, P. Gopikrishnan, B. Rosenow, L.A.N. Amaral, T. "
    "Guhr, H.E. Stanley, Random matrix approach to cross correlations "
    "in financial data, Phys. Rev. E 65 (2002) 066126.",
    "[9] M. Tumminello, T. Aste, T. Di Matteo, R.N. Mantegna, A tool "
    "for filtering information in complex systems, Proc. Natl. Acad. "
    "Sci. USA 102 (2005) 10421-10426.",
    "[10] R. Cont, Empirical properties of asset returns: stylized "
    "facts and statistical issues, Quant. Finance 1 (2001) 223-236.",
    "[11] S. Ioffe, Improved consistent sampling, weighted minhash and "
    "L1 sketching, in: Proc. IEEE Int. Conf. Data Min. (ICDM), 2010, "
    "pp. 246-255.",
    "[12] P.J. Rousseeuw, Silhouettes: a graphical aid to the "
    "interpretation and validation of cluster analysis, J. Comput. "
    "Appl. Math. 20 (1987) 53-65.",
    "[13] P. Jaccard, Etude comparative de la distribution florale dans "
    "une portion des Alpes et des Jura, Bull. Soc. Vaudoise Sci. Nat. "
    "37 (1901) 547-579.",
    "[14] K.J. Forbes, R. Rigobon, No contagion, only interdependence: "
    "measuring stock market comovements, J. Finance 57 (2002) "
    "2223-2261.",
    "[15] J.H. Ward, Hierarchical grouping to optimize an objective "
    "function, J. Am. Stat. Assoc. 58 (1963) 236-244.",
    "[16] D.H. Harbert, Weighted Jaccard decomposition detects sensor "
    "network reorganization preceding mechanical failure in turbofan "
    "engines, Scania trucks, and chemical processes, Mech. Syst. "
    "Signal Process. (2026) under review.",
    "[17] D.H. Harbert, Weighted Jaccard similarity reveals water "
    "quality parameter reorganization preceding trophic state "
    "transitions, Ecol. Indic. (2026) under review.",
    "[18] A. Tversky, Features of similarity, Psychol. Rev. 84 (1977) "
    "327-352.",
]
for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(ref)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)


# ============================================================================
# TABLES
# ============================================================================
doc.add_page_break()
add_heading(doc, "Tables")

table_caption(doc,
    "Table 1. Six structural regimes detected by the unsigned WJ "
    "rolling trajectory (all p < 0.001 by 5,000-permutation testing).")
add_word_table(doc, [
    ["Episode", "Direction", "Start", "End", "n_days", "Mean abs rho",
     "Baseline rho", "WJ", "z-score"],
    ["1", "Convergence", "2008-05-07", "2009-05-07", "504", "0.493",
     "0.309", "0.740", "-18.08"],
    ["2", "Divergence", "2016-12-07", "2017-09-08", "441", "0.176",
     "0.351", "0.681", "-22.39"],
    ["3", "Convergence", "2020-01-10", "2020-01-10", "252", "0.437",
     "0.310", "0.739", "-10.17"],
    ["4", "Convergence", "2020-05-12", "2020-08-11", "315", "0.437",
     "0.314", "0.754", "-11.68"],
    ["5", "Divergence", "2024-01-12", "2024-01-12", "252", "0.256",
     "0.351", "0.699", "-13.06"],
    ["6", "Divergence", "2024-05-14", "2024-08-14", "315", "0.244",
     "0.347", "0.690", "-16.82"],
])

table_caption(doc,
    "Table 2. Cascade stability (Kendall tau) across regime pairs by "
    "direction type. Same-direction pairs exhibit substantially "
    "stronger rank preservation than cross-direction pairs.")
add_word_table(doc, [
    ["Pair type", "n pairs", "Mean tau", "Min tau", "Max tau"],
    ["Same-direction (CONV-CONV)", "3", "0.295", "0.213", "0.525"],
    ["Same-direction (DIVE-DIVE)", "3", "0.443", "0.275", "0.664"],
    ["Cross-direction", "9", "0.119", "0.063", "0.213"],
])

table_caption(doc,
    "Table 3. WJ-native fingerprint clustering. K = 2 selected by "
    "silhouette score. Agreement with GICS taxonomy is low.")
add_word_table(doc, [
    ["Metric", "Value"],
    ["Optimal K", "2"],
    ["Silhouette score", "0.298"],
    ["Cluster A size", "319"],
    ["Cluster B size", "159"],
    ["ARI vs GICS", "0.101"],
    ["NMI vs GICS", "0.216"],
])

table_caption(doc,
    "Table 4. Pairing-family decomposition per regime: unsigned WJ, "
    "signed WJ, binary Jaccard at top-5%, Type 1 gap, Type 2 gap, and "
    "normalized Frobenius distance. Convergence regimes show large "
    "Type 1 gaps and modest Type 2 gaps; divergence regimes show the "
    "opposite. Frobenius distance is uniform across regimes.")
add_word_table(doc, [
    ["Episode", "Direction", "WJ unsigned", "WJ signed",
     "Binary Jaccard", "Type 1 gap", "Type 2 gap", "Frobenius"],
    ["1 GFC", "Convergence", "0.740", "0.918", "0.137", "+0.603",
     "+0.178", "0.097"],
    ["2 QE", "Divergence", "0.681", "0.916", "0.379", "+0.302",
     "+0.235", "0.095"],
    ["3 COVID-onset", "Convergence", "0.739", "0.919", "0.139",
     "+0.601", "+0.180", "0.098"],
    ["4 COVID-rec", "Convergence", "0.754", "0.926", "0.173", "+0.581",
     "+0.172", "0.090"],
    ["5 2024-Jan", "Divergence", "0.699", "0.919", "0.342", "+0.357",
     "+0.221", "0.096"],
    ["6 2024-H2", "Divergence", "0.690", "0.917", "0.369", "+0.321",
     "+0.227", "0.096"],
])

table_caption(doc,
    "Table 5. Implementation divergence: sign-inversion percentage and "
    "magnitude-component percentage per regime.")
add_word_table(doc, [
    ["Episode", "Direction", "Sign inversion %", "Magnitude component %"],
    ["1 GFC", "Convergence", "68.5", "31.5"],
    ["2 QE", "Divergence", "73.6", "26.4"],
    ["3 COVID-onset", "Convergence", "69.0", "31.0"],
    ["4 COVID-rec", "Convergence", "70.1", "29.9"],
    ["5 2024-Jan", "Divergence", "73.2", "26.8"],
    ["6 2024-H2", "Divergence", "73.3", "26.7"],
])

table_caption(doc,
    "Table 6. Per-window cross-trajectory Spearman rank correlations "
    "between pairing-family gaps and macro stress measures (n = 204 "
    "windows after 60-window sliding-baseline warmup).")
add_word_table(doc, [
    ["Gap metric", "vs Mean VIX", "vs Peak VIX",
     "vs Mean Realized Vol", "vs Peak Realized Vol"],
    ["Type 1 gap", "+0.685 (p<10^-29)", "+0.414 (p<10^-9)",
     "+0.727 (p<10^-34)", "+0.599 (p<10^-20)"],
    ["Type 2 gap", "+0.169 (p=0.016)", "+0.408 (p<10^-9)",
     "+0.128 (p=0.067)", "+0.231 (p<10^-3)"],
    ["Sign-inversion %", "-0.857 (p<10^-59)", "-0.639 (p<10^-24)",
     "-0.873 (p<10^-64)", "-0.793 (p<10^-44)"],
])

table_caption(doc,
    "Table 7. Permutation significance of all five signed-WJ-detected "
    "regimes (5,000 permutations of equal-day random samples from full "
    "study period). The 2022 regime is invisible to unsigned detection.")
add_word_table(doc, [
    ["Regime", "Start", "End", "n_days", "Signed WJ", "Null mean",
     "Null std", "z-score", "p-value"],
    ["GFC", "2008-07-08", "2009-05-07", "462", "0.916", "0.953",
     "0.0017", "-21.68", "<10^-3"],
    ["QE-era", "2017-01-09", "2017-09-08", "420", "0.904", "0.953",
     "0.0019", "-25.20", "<10^-3"],
    ["COVID-onset", "2020-01-10", "2020-01-10", "252", "0.919", "0.951",
     "0.0024", "-12.97", "<10^-3"],
    ["COVID-rec", "2020-05-12", "2020-08-11", "315", "0.926", "0.953",
     "0.0019", "-13.86", "<10^-3"],
    ["2022 (NEW)", "2022-07-13", "2022-08-11", "273", "0.923", "0.953",
     "0.0023", "-13.20", "<10^-3"],
])

table_caption(doc,
    "Table 8. Bootstrap 95% confidence intervals on unsigned WJ "
    "(1,000 subject-level resamples per regime) and Pearson "
    "sensitivity values.")
add_word_table(doc, [
    ["Episode", "Direction", "WJ Spearman", "95% CI Spearman",
     "WJ Pearson", "Spearman-Pearson delta"],
    ["1 GFC", "Convergence", "0.740", "[0.695, 0.779]", "0.719", "+0.022"],
    ["2 QE", "Divergence", "0.681", "[0.608, 0.737]", "0.643", "+0.038"],
    ["3 COVID-onset", "Convergence", "0.739", "[0.680, 0.762]", "0.619",
     "+0.120"],
    ["4 COVID-rec", "Convergence", "0.754", "[0.710, 0.763]", "0.652",
     "+0.102"],
    ["5 2024-Jan", "Divergence", "0.699", "[0.627, 0.726]", "0.587",
     "+0.111"],
    ["6 2024-H2", "Divergence", "0.690", "[0.622, 0.729]", "0.575",
     "+0.116"],
])


# ============================================================================
# FIGURE CAPTIONS
# ============================================================================
doc.add_page_break()
add_heading(doc, "Figure Captions")

captions = [
    ("Fig. 1.",
     "WJ rolling trajectory with data-defined regime classification "
     "(2003-2025). Top panel: each point represents a 252-trading-day "
     "window (21-day step). Colors indicate WJ-classified regime "
     "state. Red shading marks convergence episodes; blue shading "
     "marks divergence episodes. The QE-era divergence (2016-2017) "
     "produces the deepest WJ values in the dataset. Bottom panel: "
     "mean absolute Spearman correlation per window."),
    ("Fig. 2.",
     "Top 30 stocks ranked by mean fingerprint reorganization across "
     "all six detected regimes. Red bars indicate stocks in the top "
     "10% (critical tier)."),
    ("Fig. 3.",
     "Distribution of mean fingerprint reorganization across 478 "
     "stocks. Red vertical lines mark the two extreme reorganizers "
     "(CVG, EP), separated from the distribution by the largest gap "
     "(0.133)."),
    ("Fig. 4.",
     "Metric artifact test: fingerprint reorganization versus "
     "baseline mean connectivity. Dashed line shows OLS regression "
     "(R-squared = 0.437, negative slope). Red stars mark CVG and EP, "
     "both above the regression line."),
    ("Fig. 5.",
     "Directional decomposition: mean reorganization during convergence "
     "episodes versus divergence episodes for all 478 stocks. Spearman "
     "rho = 0.383. Stocks above the 1:1 line reorganize more during "
     "divergence; stocks below reorganize more during convergence."),
    ("Fig. 6.",
     "Silhouette score across K = 2-24 for WJ-native fingerprint "
     "clustering. K = 2 is optimal (silhouette = 0.298)."),
    ("Fig. 7.",
     "Hierarchical clustering dendrogram (WJ fingerprint distance, "
     "Spearman correlations)."),
    ("Fig. 8.",
     "Cascade stability heatmap (Kendall tau). Same-direction pairs "
     "(CONV-CONV, DIVE-DIVE) exhibit strong rank preservation; "
     "cross-direction pairs approach zero, demonstrating "
     "direction-dependent reorganization ordering."),
    ("Fig. 9.",
     "Pairing-family decomposition per regime. Three panels show "
     "(left) unsigned vs signed WJ side-by-side, (middle) Type 1 "
     "continuous-discrete pairing gap, and (right) Type 2 "
     "sign-treatment pairing gap. Red bars: convergence regimes; blue "
     "bars: divergence regimes. Convergence and divergence regimes "
     "show distinct pairing-family signatures."),
    ("Fig. 10.",
     "Pairing-family classification space. Each of the six structural "
     "regimes plotted by Type 1 gap (x-axis: bulk-distribution "
     "reorganization) and Type 2 gap (y-axis: sign-component "
     "reorganization). Convergence regimes (blue) and divergence "
     "regimes (red) occupy distinct quadrants."),
    ("Fig. 11.",
     "Cross-trajectory correlations between pairing-family gaps and "
     "macro stress measures. 4x3 grid: 3 gap metrics (Type 1 gap, "
     "Type 2 gap, sign-inversion percentage) versus 4 macro variables "
     "(mean VIX, peak VIX, mean realized volatility, peak realized "
     "volatility). Sign-inversion percentage shows the strongest "
     "(negative) correlation with macro stress."),
    ("Fig. 12.",
     "Per-window pairing-family trajectories with regime shading "
     "(2003-2025). Four panels: unsigned WJ, Type 1 gap, Type 2 gap, "
     "and mean VIX, all aligned to the 264 windows. Red shading: "
     "unsigned-detected convergence regimes; blue shading: "
     "unsigned-detected divergence regimes."),
]
for label, caption in captions:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(label + " ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(11)
    r1.bold = True
    r2 = p.add_run(caption)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)


ms_path = os.path.join(OUT_DIR, "Manuscript_v3.docx")
doc.save(ms_path)
print(f"Saved: {ms_path}")


# ============================================================================
# COVER LETTER (revised: softened claims)
# ============================================================================
print("Building Cover_Letter_v3.docx...")
cl = Document()
setup(cl)

add_para(cl, "Drake H. Harbert", bold=True, size=12, space_after=1,
         alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_para(cl, "Inner Architecture LLC", size=10, space_after=1,
         alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_para(cl, "Canton, OH 44721, USA", size=10, space_after=1,
         alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_para(cl, "Drake@innerarchitecturellc.com", size=10, space_after=1,
         alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_para(cl, "ORCID: 0009-0007-7740-3616", size=10, space_after=18,
         alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_para(cl, "May 8, 2026", space_after=12)
add_para(cl, "Editorial Office", space_after=1)
add_para(cl, "Physica A: Statistical Mechanics and its Applications",
         space_after=14)
add_para(cl,
    "Re: Submission of revised manuscript - 'A Pairing-Family "
    "Decomposition of the Weighted Jaccard Index Reveals Direction-"
    "Specific Reorganization Modes and Macro-Stress-Coupled Sign-"
    "Inversion Dynamics in S&P 500 Correlation Networks (2003-2025)'",
    italic=True, space_after=14)
add_para(cl, "Dear Editor,", space_after=12)
add_para(cl,
    "Please find enclosed our revised manuscript for consideration in "
    "Physica A. The revision substantially extends the original WJ-"
    "native regime detection framework with a pairing-family "
    "decomposition that separates two structurally distinct "
    "components of correlation network reorganization: magnitude "
    "(captured by unsigned WJ) and sign (captured by the gap between "
    "signed and unsigned WJ). The decomposition reveals that "
    "crisis-period reorganization is characterized by magnitude shifts "
    "while calm-era and QE-era reorganization is characterized by "
    "sign shifts. Sign-inversion percentage rank-correlates with mean "
    "realized volatility at rho = -0.873 (p < 10^-64) across 204 "
    "windows—a strong macro-stress signature for equity correlation "
    "networks.")
add_para(cl,
    "The pairing-family extension generalizes the original methodology "
    "in three respects. First, it identifies a previously unrecognized "
    "2022 structural regime (2022-07-13 to 2022-08-11; signed WJ z = "
    "-13.20, p < 10^-3 by 5,000-permutation testing) invisible to "
    "unsigned detection because the 2022 reorganization was "
    "predominantly sign-component. Second, it separates the six "
    "unsigned-detected regimes into two distinct quadrants in pairing-"
    "family space: convergence regimes (high Type 1 / low Type 2) and "
    "divergence regimes (low Type 1 / high Type 2). Third, the "
    "comparison with normalized Frobenius distance demonstrates that "
    "the regime detection is WJ-pairing-family-specific: Frobenius "
    "distance is uniform (0.090-0.098) across all six regimes despite "
    "their 8.6-fold range in WJ permutation z-scores.")
add_para(cl,
    "All previous content is preserved. The QE-era 2017 decorrelation "
    "remains the deepest unsigned WJ event in 22 years of data and is "
    "now identified as structurally distinct from the Global Financial "
    "Crisis. The directional decomposition framework, WJ-native "
    "fingerprint clustering at K = 2 (ARI = 0.101 vs GICS), and the "
    "identification of CVG and EP as extreme reorganizers across both "
    "directions are retained. Bootstrap 95% confidence intervals, "
    "Pearson correlation sensitivity, and full reproducible code "
    "(https://github.com/nwharbert8-ui/financial-wj-structural-regimes) "
    "are included.")
add_para(cl, "Highlights:")
for h in highlights:
    p = cl.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(h)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
add_para(cl, "")
add_para(cl,
    "I confirm that all listed authors have contributed to this work, "
    "that the manuscript has not been submitted elsewhere, and that no "
    "conflicts of interest exist. I look forward to the editorial "
    "review.")
add_para(cl, "Sincerely,", space_before=12, space_after=2)
add_para(cl, "", space_after=18)
add_para(cl, "Drake H. Harbert", bold=True, space_after=2)
add_para(cl, "Founder, Inner Architecture LLC", size=10, space_after=2)
add_para(cl, "Canton, OH, USA", size=10, space_after=2)
add_para(cl, "ORCID: 0009-0007-7740-3616", size=10)

cl.save(os.path.join(OUT_DIR, "Cover_Letter_v3.docx"))
print(f"Saved: Cover_Letter_v3.docx")


# ============================================================================
# HIGHLIGHTS (revised)
# ============================================================================
print("Building Highlights_v3.docx...")
hl = Document()
setup(hl)
add_para(hl, "Highlights", bold=True, size=14,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
for h in highlights:
    p = hl.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(h)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
hl.save(os.path.join(OUT_DIR, "Highlights_v3.docx"))
print(f"Saved: Highlights_v3.docx")


# ============================================================================
# DECLARATION OF INTEREST (preserve)
# ============================================================================
shutil.copy(
    os.path.join(PROJECT, "PhysicaA_Submission", "Declaration_of_Interest.docx"),
    os.path.join(OUT_DIR, "Declaration_of_Interest_v3.docx"))
print("Copied: Declaration_of_Interest_v3.docx")


# ============================================================================
# COPY FIGURES
# ============================================================================
print("Copying figures...")
src_figures = os.path.join(PROJECT, "PhysicaA_Submission", "Main_Figures")
for f in os.listdir(src_figures):
    shutil.copy(os.path.join(src_figures, f),
                 os.path.join(OUT_DIR, "Main_Figures", f))

phase2_figures_src = os.path.join(PROJECT, "results_phase2", "figures")
aug_figures_src = os.path.join(PROJECT, "results_augmentation", "figures")
fig_mapping = [
    ("pairing_decomposition.png", "Fig9.png"),
    ("triangular_regime_space.png", "Fig10.png"),
    ("macro_correlation_grid.png", "Fig11.png"),
    ("per_window_trajectories.png", "Fig12.png"),
]
for src_name, dst_name in fig_mapping:
    src1 = os.path.join(phase2_figures_src, src_name)
    src2 = os.path.join(aug_figures_src, src_name)
    src = src1 if os.path.exists(src1) else src2
    if os.path.exists(src):
        shutil.copy(src, os.path.join(OUT_DIR, "Main_Figures", dst_name))

print()
print("Build complete.")
print(f"Submission package: {OUT_DIR}")
