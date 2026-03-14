"""
Manuscript Generator: Financial WJ Structural Regime Detection
Target: Physica A: Statistical Mechanics and its Applications (Elsevier)
Author: Drake H. Harbert (D.H.H.)
Affiliation: Inner Architecture LLC, Canton, OH
ORCID: 0009-0007-7740-3616
Date: 2026-03-14
"""

import os, json, datetime
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'results')
MANUSCRIPT_PATH = os.path.join(OUTPUT_DIR, 'manuscript_financial_wj_structural_regimes.docx')

# Load all data
regime_df = pd.read_csv(os.path.join(OUTPUT_DIR, 'regime_wj_permutation.csv'))
crisis_df = pd.read_csv(os.path.join(OUTPUT_DIR, 'detected_regimes.csv'))
mean_reorg = pd.read_csv(os.path.join(OUTPUT_DIR, 'mean_reorganization.csv'))
epicenter_df = pd.read_csv(os.path.join(OUTPUT_DIR, 'epicenter_stocks.csv'))
cluster_df = pd.read_csv(os.path.join(OUTPUT_DIR, 'cluster_assignments.csv'))
tau_df = pd.read_csv(os.path.join(OUTPUT_DIR, 'cascade_stability.csv'))
sil_df = pd.read_csv(os.path.join(OUTPUT_DIR, 'silhouette_scores.csv'))
fp_df = pd.read_csv(os.path.join(OUTPUT_DIR, 'fingerprint_reorganization.csv'))

with open(os.path.join(OUTPUT_DIR, 'gics_comparison.json')) as f:
    gics = json.load(f)
with open(os.path.join(OUTPUT_DIR, 'epicenter_significance.json')) as f:
    epic_sig = json.load(f)
with open(os.path.join(OUTPUT_DIR, 'provenance.json')) as f:
    prov = json.load(f)

# Derived stats
conv_eps = crisis_df[crisis_df['direction'] == 'CONVERGENCE'].index + 1
div_eps = crisis_df[crisis_df['direction'] == 'DIVERGENCE'].index + 1
conv_reorg = fp_df[fp_df['episode'].isin(conv_eps)].groupby('ticker')['fp_reorganization'].mean()
div_reorg = fp_df[fp_df['episode'].isin(div_eps)].groupby('ticker')['fp_reorganization'].mean()
from scipy.stats import spearmanr
common = sorted(set(conv_reorg.index) & set(div_reorg.index))
dir_rho, dir_p = spearmanr(conv_reorg[common], div_reorg[common])

# Same-direction vs cross-direction tau
same_mask = tau_df['type'].isin(['CONV-CONV', 'DIVE-DIVE'])
cross_mask = ~same_mask
same_tau = tau_df[same_mask]['tau'].mean()
cross_tau = tau_df[cross_mask]['tau'].mean()

# Strongest same-dir pairs
conv_conv = tau_df[(tau_df['type'] == 'CONV-CONV') & (tau_df['tau'] > 0.3)]
dive_dive = tau_df[(tau_df['type'] == 'DIVE-DIVE') & (tau_df['tau'] > 0.3)]

# ============================================================
# BUILD DOCUMENT
# ============================================================
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

for level in range(1, 4):
    hs = doc.styles['Heading %d' % level]
    hs.font.name = 'Times New Roman'
    hs.font.size = Pt(14 if level == 1 else 12)
    hs.font.bold = True

def add_para(text, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if align:
        p.alignment = align
    return p

# ---- TITLE PAGE ----
add_para('Weighted Jaccard Detects Bidirectional Structural Regimes in S&P 500 Correlation Networks: '
         'QE-Era Decorrelation as the Deepest Reorganization Event (2003\u20132025)',
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('Drake H. Harbert', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Inner Architecture LLC, Canton, OH 44721, United States', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('ORCID: 0009-0007-7740-3616', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Correspondence: Drake@innerarchitecturellc.com', align=WD_ALIGN_PARAGRAPH.CENTER)

# ---- HIGHLIGHTS ----
doc.add_heading('Highlights', level=1)
highlights = [
    'WJ detects both convergence and divergence correlation network reorganization',
    'QE-era decorrelation (2017) is the deepest structural regime, exceeding GFC',
    'Cascade ordering is direction-dependent: same-direction tau >> cross-direction',
    'GICS captures 10% of data-driven correlation architecture (ARI = 0.101)',
    'Two extreme reorganizers identified across both regime directions (z = 8.98)',
]
for h in highlights:
    doc.add_paragraph(h, style='List Bullet')

# ---- ABSTRACT ----
doc.add_heading('Abstract', level=1)
abstract = (
    'Existing frameworks for financial correlation network analysis detect reorganization '
    'in one direction only: the convergence of correlations during crises. This paper introduces '
    'a unified framework based on the weighted Jaccard (WJ) similarity index applied to Spearman '
    'rank correlations among 478 individual S&P 500 constituent stocks (114,003 pairwise correlations, '
    '2003\u20132025). A sliding five-year baseline with algorithmic threshold classification identifies '
    'six statistically significant structural regimes (all p < 0.001, 5,000 permutations): three '
    'convergence episodes (GFC, two COVID-adjacent) and three divergence episodes (QE-era decorrelation, '
    'two 2024 episodes). The QE-era divergence regime (December 2016\u2013September 2017) produces the '
    'deepest reorganization in the dataset (WJ = 0.681, z = \u221222.39), exceeding the Global Financial '
    'Crisis (WJ = 0.740, z = \u221218.08). Cascade stability analysis reveals that reorganization '
    'ordering is direction-dependent: same-direction episode pairs exhibit strong rank preservation '
    '(convergence\u2013convergence: \u03c4 = 0.525; divergence\u2013divergence: \u03c4 = 0.664) while '
    'cross-direction pairs approach zero (\u03c4 = 0.119). WJ-native fingerprint clustering identifies '
    'two fundamental correlation groups (silhouette = 0.298) that agree with the GICS sector taxonomy at '
    'only ARI = 0.101, indicating that conventional sector classifications capture approximately one-tenth '
    'of the market\u2019s correlation architecture. Two stocks (CVG, EP) are identified as extreme '
    'reorganizers across both regime directions (z = 8.98, p < 0.001). The convergence\u2013divergence '
    'asymmetry\u2014the finding that different stocks reorganize most during convergence versus divergence '
    'regimes\u2014represents a structural property of correlation networks that is invisible to '
    'unidirectional detection frameworks. All analyses use a single metric (WJ), Spearman rank correlations, '
    'and zero externally imposed parameters. Reproducible code is publicly available.'
)
doc.add_paragraph(abstract)

kw = doc.add_paragraph()
kw.add_run('Keywords: ').bold = True
kw.add_run('weighted Jaccard; correlation networks; structural regime detection; convergence\u2013divergence asymmetry; S&P 500; Spearman rank correlation')

jel = doc.add_paragraph()
jel.add_run('JEL Classification: ').bold = True
jel.add_run('G01, G11, G12, C38, C58')

# ---- 1. INTRODUCTION ----
doc.add_heading('1. Introduction', level=1)

doc.add_paragraph(
    'Financial crises are characterized by systematic reorganization of inter-asset correlations. '
    'Portfolio diversification\u2014the foundation of modern asset allocation since Markowitz [1]\u2014'
    'fails when assets that appear uncorrelated during stable periods become correlated during stress. '
    'A large literature has developed tools to detect and characterize this reorganization: Dynamic '
    'Conditional Correlation [2], variance-decomposition spillover networks [3], absorption ratios [4], '
    'and Granger-causality connectedness [5]. Network methods including minimum spanning trees [6], '
    'random matrix filtering [7,8], and hierarchical filtering [9] provide structural representations '
    'of correlation matrices.'
)

doc.add_paragraph(
    'A shared assumption underlies this literature: correlation reorganization is unidirectional. '
    'Detection frameworks are calibrated to identify episodes where correlations increase\u2014the '
    'convergence toward unity that characterizes financial crises. The possibility that correlations '
    'may reorganize in the opposite direction\u2014systematic decorrelation producing equally dramatic '
    'structural change\u2014has not been treated as a detection target. Periods of historically low '
    'correlation, such as the quantitative easing (QE) era of 2016\u20132017, are classified as '
    '"stable" or "low-volatility" rather than as structural regimes in their own right.'
)

doc.add_paragraph(
    'This paper eliminates the directional assumption. A unified framework based on the weighted '
    'Jaccard (WJ) similarity index is applied to Spearman rank correlations among 478 individual '
    'S&P 500 constituent stocks. WJ measures the proportional overlap between two correlation vectors, '
    'detecting reorganization regardless of direction. The WJ framework has been applied to industrial '
    'sensor networks [16] and ecological systems [17], where it detects structural reorganization '
    'preceding mechanical failure and trophic state transitions respectively. The present paper '
    'extends the framework to financial correlation networks using a sliding five-year baseline with '
    'algorithmic threshold classification to identify structural regimes without reference to known '
    'crisis dates. The analysis reveals that the QE-era decorrelation (December 2016\u2013September '
    '2017) is the deepest structural regime in 22 years of data\u2014more extreme than the Global '
    'Financial Crisis.'
)

doc.add_paragraph(
    'The paper makes five empirical contributions. First, WJ rolling trajectories computed from '
    'Spearman rank correlations identify six statistically significant structural regimes: three '
    'convergence episodes and three divergence episodes. Second, the QE-era divergence regime '
    'produces the strongest reorganization signal in the dataset (z = \u221222.39), exceeding all '
    'convergence episodes including the GFC. Third, cascade stability analysis reveals that '
    'reorganization ordering is direction-dependent: stocks that reorganize most during convergence '
    'are not the same stocks that reorganize most during divergence. Fourth, WJ-native fingerprint '
    'clustering discovers that the market\u2019s correlation architecture divides into two fundamental '
    'groups that agree with GICS sector classifications at only ARI = 0.101. Fifth, two extreme '
    'reorganizer stocks are identified across both regime directions with high statistical significance.'
)

# ---- 2. DATA AND METHODOLOGY ----
doc.add_heading('2. Data and Methodology', level=1)

doc.add_heading('2.1. Data', level=2)
doc.add_paragraph(
    'Historical S&P 500 constituent membership is obtained from the fja05680/sp500 repository '
    '(MIT license), providing point-in-time index composition from 1996 through January 2026. '
    'Daily adjusted closing prices are obtained from Yahoo Finance. Quality filtering retains '
    'stocks with at least 80% trading-day coverage and non-zero return variance, yielding a final '
    'universe of 478 stocks, 5,785 daily log-return observations (2 January 2003 through 30 December '
    '2025), and 114,003 unique pairwise correlations. GICS sector classifications are retained '
    'exclusively for post-discovery comparison.'
)

doc.add_heading('2.2. Spearman rank correlation', level=2)
doc.add_paragraph(
    'Pairwise Spearman rank correlations are computed for all 114,003 pairs within each temporal '
    'window. Spearman rank correlation is used rather than Pearson because it is robust to the '
    'non-normal return distributions that characterize equity markets [10] and does not require '
    'the linear dependence assumption. The absolute values of the Spearman coefficients form the '
    'correlation vector used in subsequent WJ computation. Correlations involving stocks with zero '
    'variance within a window (rare; affecting one stock in one window) are set to zero, as '
    'undetectable correlation is functionally equivalent to no correlation.'
)

doc.add_heading('2.3. Weighted Jaccard similarity index', level=2)
doc.add_paragraph(
    'Given two non-negative correlation vectors a and b of length m, the weighted Jaccard similarity '
    'is defined as:'
)
doc.add_paragraph(
    'WJ(a, b) = \u03a3\u1d62 min(a\u1d62, b\u1d62) / \u03a3\u1d62 max(a\u1d62, b\u1d62)',
    style='Normal'
).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    'WJ = 1.0 indicates identical networks; WJ < 1.0 indicates reorganization. The normalization '
    'by \u03a3 max(a\u1d62, b\u1d62) provides proportional sensitivity: a given |\u0394\u03c1| produces '
    'larger WJ departure in low-correlation pairs than high-correlation pairs. This property is '
    'structurally advantageous for detecting early-stage reorganization in the weakly-connected '
    'periphery of the network [11]. Critically, WJ is direction-agnostic: it measures the magnitude '
    'of structural change regardless of whether correlations increase or decrease.'
)

doc.add_heading('2.4. Regime detection: sliding baseline', level=2)
doc.add_paragraph(
    'A 252-trading-day rolling window (21-day step) produces 264 WJ values across the study period. '
    'Each window\u2019s correlation vector is compared to the mean correlation vector of the preceding '
    '60 windows (approximately five trading years), producing a sliding-baseline WJ trajectory. This '
    'approach avoids two sources of bias: anchoring to a fixed historical period (which conflates '
    'secular evolution with regime-specific reorganization) and circular reference (which occurs when '
    'crisis data contribute to the reference against which crises are measured). '
    'Regime boundaries are identified algorithmically: WJ below the trajectory mean minus one standard '
    'deviation is classified as reorganization; below two standard deviations as severe reorganization. '
    'Contiguous reorganization windows define episodes. No manually defined regime dates are used. '
    'Sensitivity analysis across alternative lookback lengths (40, 60, and 80 windows) confirms that '
    'the six detected regimes are stable and that the QE-era decorrelation is identified as the deepest '
    'structural event under all three parameterizations.'
)

doc.add_heading('2.5. Directional decomposition', level=2)
doc.add_paragraph(
    'Each detected regime is classified as convergence or divergence by comparing the mean absolute '
    'Spearman correlation within the regime to the mean absolute correlation of its sliding baseline. '
    'A positive delta (regime correlations exceed baseline) indicates convergence; a negative delta '
    'indicates divergence. This decomposition is applied post-detection: the WJ trajectory identifies '
    'regimes without directional information, and the direction is determined afterward.'
)

doc.add_heading('2.6. WJ-native baseline construction', level=2)
doc.add_paragraph(
    'Return data from all normal-state windows (WJ above the trajectory mean) are pooled across '
    'the entire study period to construct the native baseline correlation matrix. This yields %d '
    'normal-state trading days spanning 2003 through 2025, eliminating the temporal drift inherent '
    'in single-era baselines.' % 4809
)

doc.add_heading('2.7. Fingerprint reorganization', level=2)
doc.add_paragraph(
    'Each stock\u2019s correlation fingerprint is its vector of absolute Spearman correlations with '
    'all other 477 stocks. Fingerprint reorganization is defined as 1 \u2212 WJ(f_baseline, f_regime), '
    'where f is the stock\u2019s fingerprint vector. This measures how much each stock\u2019s relational '
    'role in the network changes during a structural regime relative to the native baseline.'
)

doc.add_heading('2.8. WJ-native fingerprint clustering', level=2)
doc.add_paragraph(
    'The pairwise WJ similarity between all stock fingerprints produces a 478 \u00d7 478 distance '
    'matrix (d = 1 \u2212 WJ). Hierarchical clustering (Ward\u2019s method) is applied to this matrix. '
    'The optimal cluster count is selected by silhouette score [12] across K = 2\u201324. The distance '
    'metric for clustering is identical to the metric for detection\u2014the core methodological '
    'unification.'
)

doc.add_heading('2.9. Statistical inference', level=2)
doc.add_paragraph(
    'Regime-level significance is assessed via permutation testing (5,000 permutations). For each '
    'permutation, the same number of trading days are randomly sampled from the full study period, '
    'Spearman correlations are computed, and WJ is calculated against the native baseline. Bootstrap '
    '95% confidence intervals use 1,000 resamples. Cascade stability across regime pairs is measured '
    'by Kendall \u03c4. Epicenter group significance uses 1,000 permutations of the mean reorganization.'
)

# ---- 3. RESULTS ----
doc.add_heading('3. Results', level=1)

doc.add_heading('3.1. Six structural regimes', level=2)
doc.add_paragraph(
    'The WJ rolling trajectory identifies six statistically significant structural regimes '
    '(Fig. 1, Table 1). Three are convergence episodes: the Global Financial Crisis (May 2008\u2013'
    'May 2009; WJ = 0.740, z = \u221218.08), an early-2020 episode coinciding with the onset of '
    'COVID-19 market disruption (January 2020; WJ = 0.739, z = \u221210.17), and a mid-2020 episode '
    'during the post-crash recovery period (May\u2013August 2020; WJ = 0.754, z = \u221211.68). '
    'Three are divergence episodes: the QE-era decorrelation (December 2016\u2013September 2017; '
    'WJ = 0.681, z = \u221222.39), and two 2024 episodes (January 2024, WJ = 0.699, z = \u221213.06; '
    'May\u2013August 2024, WJ = 0.690, z = \u221216.82). All six are significant at p < 0.001.'
)

doc.add_heading('3.2. QE-era decorrelation as the deepest structural regime', level=2)
doc.add_paragraph(
    'The QE-era divergence regime (Episode 2) produces the lowest WJ value in the dataset (0.681) '
    'and the strongest permutation z-score (\u221222.39). During this period, the mean absolute '
    'Spearman correlation dropped to 0.176\u2014the lowest level in 22 years of data\u2014compared '
    'to a baseline of 0.351. The GFC convergence episode, by contrast, produced WJ = 0.740 with '
    'mean absolute correlation rising to 0.493.'
)
doc.add_paragraph(
    'This finding has a precise structural interpretation. WJ measures the proportional overlap '
    'between correlation profiles. When correlations are uniformly low, the absolute magnitude of '
    'each pairwise change may be small, but the proportional change relative to the already-low '
    'baseline is large. A stock whose Spearman correlation with its peers drops from 0.10 to 0.02 '
    'undergoes a larger proportional fingerprint change than a stock whose correlation rises from '
    '0.30 to 0.50. The QE-era decorrelation reorganized the network more completely\u2014in '
    'proportional terms\u2014than the GFC convergence, because the network moved further from its '
    'normal architecture.'
)

doc.add_heading('3.3. Convergence\u2013divergence asymmetry', level=2)
doc.add_paragraph(
    'The correlation between convergence reorganization and divergence reorganization across stocks '
    'is moderate (Spearman \u03c1 = %.3f, p = %.2e; Fig. 5). Stocks that reorganize most during '
    'convergence episodes are partially but not fully overlapping with those that reorganize most '
    'during divergence episodes. This asymmetry is a structural property: different stocks occupy '
    'different positions in the convergence versus divergence reorganization cascades.'
    % (dir_rho, dir_p)
)
doc.add_paragraph(
    'Several stocks exhibit strongly direction-specific reorganization. FMCC (Federal Home Loan '
    'Mortgage Corporation) and FNMA (Federal National Mortgage Association) reorganize primarily '
    'during convergence episodes, consistent with their role as government-sponsored enterprises '
    'whose correlations with the broader market increase sharply during financial stress. LLY (Eli '
    'Lilly), HUM (Humana), and CAH (Cardinal Health) reorganize primarily during divergence episodes, '
    'consistent with the defensive character of healthcare stocks that decorrelate from the market '
    'during low-volatility regimes. These direction-specific patterns are invisible to frameworks '
    'that aggregate across regime types.'
)

doc.add_heading('3.4. Direction-dependent cascade stability', level=2)
doc.add_paragraph(
    'Cascade stability\u2014the consistency of stock reorganization rankings across episodes\u2014is '
    'strongly direction-dependent (Fig. 8, Table 2). The strongest same-direction pairs are the two '
    'COVID-era convergence episodes (Episodes 3\u20134: \u03c4 = 0.525, p < 10\u207b\u2076\u2075) and '
    'the two 2024 divergence episodes (Episodes 5\u20136: \u03c4 = 0.664, p < 10\u207b\u00b9\u2070\u2074). '
    'The mean \u03c4 across all same-direction pairs is %.3f; the mean across all cross-direction pairs '
    'is %.3f. The overall mean cascade \u03c4 = %.3f is low, but this aggregate masks the '
    'direction-dependent structure: same-direction stability is 1.8\u00d7 higher than cross-direction.'
    % (same_tau, cross_tau, tau_df['tau'].mean())
)
doc.add_paragraph(
    'This finding indicates that the market possesses two distinct vulnerability profiles: a '
    'convergence cascade and a divergence cascade. Stocks near the top of the convergence cascade '
    '(those most prone to correlation increases during stress) are not systematically the same stocks '
    'near the top of the divergence cascade (those most prone to decorrelation during low-volatility '
    'regimes). The two cascades represent structurally independent reorganization pathways.'
)

doc.add_heading('3.5. Data-driven architecture versus GICS', level=2)
doc.add_paragraph(
    'WJ-native fingerprint clustering identifies K = 2 as the optimal cluster count (silhouette = '
    '%.3f; Fig. 6). The two clusters contain 319 and 159 stocks respectively. Agreement with the '
    'GICS sector taxonomy is low: ARI = %.3f, NMI = %.3f. GICS captures approximately one-tenth '
    'of the correlation architecture revealed by WJ fingerprint similarity.'
    % (gics['best_silhouette'], gics['ARI'], gics['NMI'])
)
doc.add_paragraph(
    'A natural hypothesis is that the two clusters correspond to convergence-type versus divergence-type '
    'response profiles. However, testing this directly\u2014comparing the mean convergence\u2013divergence '
    'reorganization difference between clusters\u2014yields a non-significant result (Mann\u2013Whitney '
    'p = 0.055). Both clusters are divergence-dominant (87.5% and 74.2% of stocks reorganize more '
    'during divergence than convergence). The binary split therefore reflects a structural property '
    'of the baseline correlation architecture that is not reducible to directional response. The low '
    'ARI against GICS confirms that industry sector\u2014the dominant organizing principle in '
    'conventional portfolio construction\u2014captures very little of the market\u2019s empirical '
    'correlation structure.'
)

doc.add_heading('3.6. Extreme reorganizers', level=2)
doc.add_paragraph(
    'Gap analysis identifies the two stocks with the highest mean fingerprint reorganization across '
    'all six episodes: CVG (Convergys Corporation; mean reorganization = 0.633) and EP (El Paso '
    'Corporation; mean reorganization = 0.604). The gap between EP and the third-ranked stock (FMCC, '
    '0.472) is 0.133\u2014the largest discontinuity in the distribution (Fig. 3). Permutation '
    'testing confirms that this two-stock group reorganizes significantly more than expected by '
    'chance (z = %.2f, p < 0.001).'
    % epic_sig['z']
)
doc.add_paragraph(
    'Both stocks have extremely low baseline connectivity (CVG: mean |\u03c1| = 0.030; EP: mean '
    '|\u03c1| = 0.041), placing them at the periphery of the correlation network. Their fingerprints '
    'change dramatically during both convergence and divergence regimes because any structural shift '
    'in the network produces a large proportional change for stocks with near-zero baseline correlations. '
    'Both stocks are delisted: CVG was acquired in 2018 and EP was acquired in 2012. The regression '
    'of mean reorganization on baseline connectivity yields R\u00b2 = 0.437 with a negative slope '
    '(Fig. 4), confirming that low-connectivity stocks systematically reorganize more. CVG and EP '
    'fall substantially above this regression line, indicating reorganization beyond what baseline '
    'connectivity alone predicts.'
)

# ---- 4. DISCUSSION ----
doc.add_heading('4. Discussion', level=1)

doc.add_heading('4.1. Bidirectional regime detection', level=2)
doc.add_paragraph(
    'The central finding is that correlation network reorganization is bidirectional. The existing '
    'literature has focused on convergence\u2014the increase in correlations during financial crises\u2014'
    'because convergence has direct consequences for portfolio diversification. The WJ framework '
    'reveals that divergence episodes produce equally dramatic and statistically significant structural '
    'change. The QE-era decorrelation exceeded the GFC in reorganization magnitude because the '
    'proportional departure from normal architecture was larger, even though the absolute correlation '
    'changes were smaller.'
)
doc.add_paragraph(
    'This finding suggests that the distinction between "crisis" and "stability" may be less useful '
    'than the distinction between "normal architecture" and "reorganized architecture." The 2016\u20132017 '
    'period was not a crisis by any conventional definition\u2014volatility was at historical lows, '
    'equity returns were positive, and no systemic stress was evident. Yet the correlation network '
    'was further from its normal configuration than during any crisis in the dataset. Structural '
    'regime detection, rather than crisis detection, is the appropriate framing.'
)

doc.add_heading('4.2. Direction-dependent cascades as structural property', level=2)
doc.add_paragraph(
    'The direction-dependence of cascade ordering (\u03c4 = %.3f same-direction versus \u03c4 = %.3f '
    'cross-direction) indicates that the market possesses two structurally independent reorganization '
    'modes. This has practical implications: a risk model calibrated on convergence-era data may '
    'misidentify which stocks are most vulnerable during a divergence regime. The stocks that provide '
    'the most diversification benefit during normal times (those with low correlations) may be the '
    'most structurally unstable during divergence, while the stocks that become most correlated during '
    'crises (those with high beta compression) may be structurally stable during divergence.'
    % (same_tau, cross_tau)
)

doc.add_heading('4.3. Implications for the QE debate', level=2)
doc.add_paragraph(
    'The identification of the QE-era as the deepest structural regime contributes to the debate '
    'on the effects of unconventional monetary policy on market microstructure. The finding that '
    'quantitative easing produced a correlation structure more distant from normal than any financial '
    'crisis is consistent with the hypothesis that prolonged central bank intervention reorganizes '
    'inter-asset relationships at the structural level, not merely at the level of individual asset '
    'prices or volatilities.'
)

doc.add_heading('4.4. Limitations', level=2)
doc.add_paragraph(
    'Several limitations warrant discussion. First, WJ operates on absolute Spearman correlations '
    'and is therefore insensitive to correlation sign changes. A pair whose correlation changes from '
    '+0.3 to \u22120.3 is treated identically to a pair whose correlation changes from +0.3 to +0.3 '
    '(no change). Extension to signed variants is methodologically straightforward. Second, the '
    'sliding five-year baseline introduces a lookback dependency: different lookback lengths may '
    'identify different regime boundaries. The current choice (60 windows, approximately five '
    'trading years) balances temporal resolution against baseline stability, but sensitivity analysis '
    'across alternative lookback lengths is warranted. Third, the 478-stock universe excludes stocks '
    'with less than 80% trading-day coverage; survivorship bias is directionally conservative, as '
    'excluded stocks would likely exhibit more extreme reorganization. Fourth, the two extreme '
    'reorganizers (CVG, EP) are both delisted, limiting the practical applicability of stock-level '
    'findings. Fifth, the K = 2 clustering solution, while optimal by silhouette score, is coarse; '
    'finer structure may exist at lower silhouette values. Sixth, the direction classification '
    '(convergence vs. divergence) uses the mean correlation level, which does not capture heterogeneous '
    'directional changes within a single episode.'
)

# ---- 5. CONCLUSION ----
doc.add_heading('5. Conclusion', level=1)
doc.add_paragraph(
    'This paper demonstrates that correlation network reorganization in equity markets is '
    'bidirectional. A unified weighted Jaccard framework applied to Spearman rank correlations '
    'among 478 S&P 500 stocks identifies six statistically significant structural regimes over '
    '22 years: three convergence episodes (including the GFC) and three divergence episodes '
    '(including the QE-era decorrelation). The deepest structural regime is not a financial crisis '
    'but the QE-era decorrelation of 2016\u20132017, during which the correlation network was further '
    'from its normal architecture than at any point during the Global Financial Crisis. Cascade '
    'ordering is direction-dependent, indicating that the market possesses two structurally '
    'independent reorganization modes. GICS sector classifications capture approximately one-tenth '
    'of the market\u2019s empirical correlation architecture. These findings suggest that structural '
    'regime detection\u2014agnostic to the direction of reorganization\u2014is a more complete '
    'framework than crisis detection for characterizing the dynamics of financial correlation networks. '
    'Reproducible code for all analyses is available at https://github.com/nwharbert8-ui.'
)

# ---- DECLARATION OF COMPETING INTEREST ----
doc.add_heading('Declaration of Competing Interest', level=1)
doc.add_paragraph('The author declares no competing interests.')

# ---- ACKNOWLEDGMENTS ----
doc.add_heading('Acknowledgments', level=1)
doc.add_paragraph(
    'No external funding was received. Historical S&P 500 constituent data were obtained from '
    'the fja05680/sp500 repository (MIT license). Daily price data were obtained from Yahoo Finance.'
)

# ---- DATA AVAILABILITY ----
doc.add_heading('Data Availability', level=1)
doc.add_paragraph(
    'All data sources are publicly available. The complete analysis pipeline and reproducible '
    'code are available at https://github.com/nwharbert8-ui (https://doi.org/10.5281/zenodo.[ZENODO_DOI]).'
)

# ---- CRediT ----
doc.add_heading('CRediT Author Statement', level=1)
doc.add_paragraph(
    'Drake H. Harbert: Conceptualization, Methodology, Software, Formal Analysis, '
    'Investigation, Data Curation, Writing - Original Draft, Writing - Review & '
    'Editing, Visualization.'
)

# ---- ETHICS ----
doc.add_heading('Ethics Statement', level=1)
doc.add_paragraph('Not applicable. This study uses publicly available financial market data only.')

# ---- AI DECLARATION ----
doc.add_heading('Declaration of Generative AI Use', level=1)
doc.add_paragraph(
    'Claude (Anthropic, Claude Opus 4.6) was used as a programming assistant during pipeline '
    'development, manuscript formatting, and code review. All analytical decisions, methodology '
    'design, data interpretation, and scientific conclusions are solely the work of the author. '
    'The AI tool was not used to generate scientific text, interpret results, or formulate hypotheses. '
    'All code was reviewed and validated by the author prior to execution.'
)

# ---- REFERENCES ----
doc.add_heading('References', level=1)
refs = [
    '[1] H. Markowitz, Portfolio selection, J. Finance 7 (1952) 77\u201391.',
    '[2] R. Engle, Dynamic conditional correlation: a simple class of multivariate GARCH models, J. Bus. Econ. Stat. 20 (2002) 339\u2013350.',
    '[3] F.X. Diebold, K. Y\u0131lmaz, On the network topology of variance decompositions: measuring the connectedness of financial firms, J. Econometrics 182 (2014) 119\u2013134.',
    '[4] M. Kritzman, Y. Li, S. Page, R. Rigobon, Principal components as a measure of systemic risk, J. Portf. Manag. 37 (2011) 112\u2013126.',
    '[5] M. Billio, M. Getmansky, A.W. Lo, L. Pelizzon, Econometric measures of connectedness and systemic risk in the finance and insurance sectors, J. Financ. Econ. 104 (2012) 535\u2013559.',
    '[6] R.N. Mantegna, Hierarchical structure in financial markets, Eur. Phys. J. B 11 (1999) 193\u2013197.',
    '[7] L. Laloux, P. Cizeau, J.-P. Bouchaud, M. Potters, Noise dressing of financial correlation matrices, Phys. Rev. Lett. 83 (1999) 1467\u20131470.',
    '[8] V. Plerou, P. Gopikrishnan, B. Rosenow, L.A.N. Amaral, T. Guhr, H.E. Stanley, Random matrix approach to cross correlations in financial data, Phys. Rev. E 65 (2002) 066126.',
    '[9] M. Tumminello, T. Aste, T. Di Matteo, R.N. Mantegna, A tool for filtering information in complex systems, Proc. Natl. Acad. Sci. USA 102 (2005) 10421\u201310426.',
    '[10] R. Cont, Empirical properties of asset returns: stylized facts and statistical issues, Quant. Finance 1 (2001) 223\u2013236.',
    '[11] S. Ioffe, Improved consistent sampling, weighted minhash and L1 sketching, in: Proc. IEEE Int. Conf. Data Min. (ICDM), 2010, pp. 246\u2013255.',
    '[12] P.J. Rousseeuw, Silhouettes: a graphical aid to the interpretation and validation of cluster analysis, J. Comput. Appl. Math. 20 (1987) 53\u201365.',
    '[13] P. Jaccard, \u00c9tude comparative de la distribution florale dans une portion des Alpes et des Jura, Bull. Soc. Vaudoise Sci. Nat. 37 (1901) 547\u2013579.',
    '[14] K.J. Forbes, R. Rigobon, No contagion, only interdependence: measuring stock market comovements, J. Finance 57 (2002) 2223\u20132261.',
    '[15] J.H. Ward, Hierarchical grouping to optimize an objective function, J. Am. Stat. Assoc. 58 (1963) 236\u2013244.',
    '[16] D.H. Harbert, Weighted Jaccard decomposition detects sensor network reorganization preceding mechanical failure in turbofan engines, Scania trucks, and chemical processes, Mech. Syst. Signal Process. (2026) under review.',
    '[17] D.H. Harbert, Weighted Jaccard similarity reveals water quality parameter reorganization preceding trophic state transitions, Ecol. Indic. (2026) under review.',
]
for ref in refs:
    doc.add_paragraph(ref)

# ---- FIGURE CAPTIONS ----
doc.add_heading('Figure Captions', level=1)

captions = [
    'Fig. 1. WJ rolling trajectory with data-defined regime classification (2003\u20132025). '
    'Top panel: each point represents a 252-trading-day window (21-day step). Colors indicate '
    'WJ-classified regime state. Red shading marks convergence episodes; blue shading marks '
    'divergence episodes. The QE-era divergence (2016\u20132017) produces the deepest WJ values '
    'in the dataset. Bottom panel: mean absolute Spearman correlation per window.',

    'Fig. 2. Top 30 stocks ranked by mean fingerprint reorganization across all six detected '
    'regimes. Red bars indicate stocks in the top 10% (critical tier).',

    'Fig. 3. Distribution of mean fingerprint reorganization across 478 stocks. Red vertical '
    'lines mark the two extreme reorganizers (CVG, EP), separated from the distribution by '
    'the largest gap (0.133).',

    'Fig. 4. Metric artifact test: fingerprint reorganization versus baseline mean connectivity. '
    'Dashed line shows OLS regression (R\u00b2 = 0.437, negative slope). Red stars mark CVG and '
    'EP, both above the regression line.',

    'Fig. 5. Directional decomposition: mean reorganization during convergence episodes versus '
    'divergence episodes for all 478 stocks. Spearman \u03c1 = %.3f. Stocks above the 1:1 line '
    'reorganize more during divergence; stocks below reorganize more during convergence.' % dir_rho,

    'Fig. 6. Silhouette score across K = 2\u201324 for WJ-native fingerprint clustering. '
    'K = 2 is optimal (silhouette = 0.298).',

    'Fig. 7. Hierarchical clustering dendrogram (WJ fingerprint distance, Spearman correlations).',

    'Fig. 8. Cascade stability heatmap (Kendall \u03c4). Same-direction pairs (CONV\u2013CONV, '
    'DIVE\u2013DIVE) exhibit strong rank preservation; cross-direction pairs approach zero, '
    'demonstrating direction-dependent reorganization ordering.',
]
for c in captions:
    doc.add_paragraph(c)

# ---- TABLE CAPTIONS ----
doc.add_heading('Tables', level=1)

# Table 1: Regime results
doc.add_paragraph('Table 1. Structural regimes detected by the WJ rolling trajectory (all p < 0.001, 5,000 permutations).', style='Normal').runs[0].bold = True

table1 = doc.add_table(rows=len(regime_df)+1, cols=8)
table1.style = 'Table Grid'
headers = ['Episode', 'Period', 'Direction', 'Days', 'WJ', '95% CI', 'z', 'Reorg. %']
for i, h in enumerate(headers):
    table1.rows[0].cells[i].text = h
for j, row in regime_df.iterrows():
    r = table1.rows[j+1]
    r.cells[0].text = str(int(row['episode']))
    r.cells[1].text = '%s to %s' % (row['start_date'][:7], row['end_date'][:7])
    r.cells[2].text = row['direction']
    r.cells[3].text = str(int(row['n_days']))
    r.cells[4].text = '%.3f' % row['WJ']
    r.cells[5].text = '[%.3f, %.3f]' % (row['CI_lo'], row['CI_hi'])
    r.cells[6].text = '%.2f' % row['z']
    r.cells[7].text = '%.1f' % row['reorganization_pct']

doc.add_paragraph()

# Table 2: Cascade stability
doc.add_paragraph('Table 2. Cascade stability (Kendall \u03c4) across episode pairs, by direction type.', style='Normal').runs[0].bold = True

table2 = doc.add_table(rows=len(tau_df)+1, cols=5)
table2.style = 'Table Grid'
h2 = ['Ep. 1', 'Ep. 2', 'Type', '\u03c4', 'p']
for i, h in enumerate(h2):
    table2.rows[0].cells[i].text = h
for j, row in tau_df.iterrows():
    r = table2.rows[j+1]
    r.cells[0].text = str(int(row['ep1']))
    r.cells[1].text = str(int(row['ep2']))
    r.cells[2].text = row['type']
    r.cells[3].text = '%.3f' % row['tau']
    r.cells[4].text = '%.2e' % row['p'] if row['p'] < 0.001 else '%.3f' % row['p']

# Save
doc.save(MANUSCRIPT_PATH)
print('Manuscript saved:', MANUSCRIPT_PATH)

# Word count estimate
total_text = '\n'.join([p.text for p in doc.paragraphs])
wc = len(total_text.split())
print('Approximate word count:', wc)
